"""Convert the FAQ markdown draft to a Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.font.name = "Calibri"

with open("/home/user/sentinel-website/docs/faq-draft.md", "r") as f:
    content = f.read()

lines = content.split("\n")
i = 0
skip_toc = False

def add_table(header_line, rows):
    """Add a formatted table to the document."""
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_text in rows:
        cells_text = [c.strip() for c in row_text.strip("|").split("|")]
        row = table.add_row()
        for j, ct in enumerate(cells_text):
            if j < num_cols:
                row.cells[j].text = ct
                for paragraph in row.cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph()  # spacing after table


def process_inline(paragraph, text):
    """Process bold, italic, and links in text and add runs to a paragraph."""
    # Pattern for **bold**, *italic*, [link](url)
    pattern = re.compile(
        r"\*\*\[([^\]]+)\]\*\*"  # **[Tag]** -> bold tag
        r"|\*\*([^*]+)\*\*"      # **bold**
        r"|\*([^*]+)\*"          # *italic*
        r"|\[([^\]]+)\]\(([^)]+)\)"  # [text](url)
        r"|`([^`]+)`"            # `code`
    )
    last_end = 0
    for m in pattern.finditer(text):
        # Add text before match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(1):  # **[Tag]**
            run = paragraph.add_run(f"[{m.group(1)}]")
            run.bold = True
        elif m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # [text](url)
            run = paragraph.add_run(m.group(4))
            run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
            run.underline = True
        elif m.group(6):  # `code`
            run = paragraph.add_run(m.group(6))
            run.font.name = "Consolas"
            run.font.size = Pt(10)

        last_end = m.end()

    # Remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


while i < len(lines):
    line = lines[i]

    # Skip the YAML/front-matter style lines
    if line.strip() == "---":
        i += 1
        continue

    # Skip TOC section
    if line.strip() == "## Table of Contents":
        skip_toc = True
        i += 1
        continue
    if skip_toc:
        if line.startswith("## ") and line.strip() != "## Table of Contents":
            skip_toc = False
        else:
            i += 1
            continue

    # Title
    if line.startswith("# ") and not line.startswith("## "):
        title = line[2:].strip()
        p = doc.add_heading(title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Blockquote (subtitle / date)
    if line.startswith("> "):
        text = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        i += 1
        continue

    # H2
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue

    # H3 (questions)
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue

    # Table detection
    if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i + 1]):
        header_line = line
        i += 2  # skip header and separator
        data_rows = []
        while i < len(lines) and "|" in lines[i] and lines[i].strip():
            data_rows.append(lines[i])
            i += 1
        add_table(header_line, data_rows)
        continue

    # Bullet points
    if line.startswith("- **"):
        # Bold-lead bullet
        p = doc.add_paragraph(style="List Bullet")
        match = re.match(r"^- \*\*([^*]+)\*\*:?\s*(.*)", line)
        if match:
            run = p.add_run(match.group(1))
            run.bold = True
            if match.group(2):
                p.add_run(": " + match.group(2) if not match.group(2).startswith(":") else match.group(2))
        else:
            process_inline(p, line[2:])
        i += 1
        continue

    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        process_inline(p, line[2:].strip())
        i += 1
        continue

    # Regular paragraph
    text = line.strip()
    if text:
        p = doc.add_paragraph()
        process_inline(p, text)

    i += 1

output_path = "/home/user/sentinel-website/docs/faq-draft.docx"
doc.save(output_path)
print(f"Word document saved to {output_path}")
